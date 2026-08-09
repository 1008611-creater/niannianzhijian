from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"E:\codex\niannianai\niannianai")
OUT_DIR = ROOT / "documents" / "authoritative-docs" / "niannian-ai-web"
OUT = OUT_DIR / "念念AI主站权威项目文档V2.docx"
IDEAL = ROOT / "documents" / "authoritative-docs" / "visuals" / "niannian-ai-web_ideal-home.png"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=(255,255,255), size=8)
        set_cell_shading(table.rows[0].cells[i], "1F2937")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], "F3F4F6")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p

def add_stage(doc, name, goal, inputs, deliverable, plan, main_skill, support_skill, agent, boundary, acceptance, next_input, prompt):
    doc.add_heading(name, level=2)
    for label, value in [
        ("阶段目标", goal), ("当前真实输入", inputs), ("唯一交付物", deliverable),
        ("可执行方案", plan), ("唯一主 Skill", main_skill), ("直接支持 Skill", support_skill),
        ("Agent / AGENTS.md 内容", agent), ("不可越过的产品边界", boundary),
        ("验收标准", acceptance), ("下一阶段输入", next_input),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + "：")
        r.bold = True
        p.add_run(value)
    p = doc.add_paragraph()
    r = p.add_run("可复制给 Codex 的阶段提示词：")
    r.bold = True
    p.add_run("")
    add_code(doc, prompt)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(9.5)
for style_name, size, color in [("Title", 24, (17,24,39)), ("Heading 1", 16, (17,24,39)), ("Heading 2", 12, (31,41,55))]:
    st = styles[style_name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(*color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("念念 AI 主站权威项目文档 V2")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(17,24,39)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("前期产品与工程唯一入口 · 事实优先 · 不发布版本").italic = True
doc.add_paragraph("版本：V2 | 日期：2026-08-10 | 项目：念念 AI（AI 短剧创作平台）")
doc.add_paragraph("文档状态：可审阅前期基线。任何未被当前文件、线上回读或截图证据证明的内容均标为“待核实”，本文件不授权线上改动或发布。")

doc.add_heading("1. 定位与真实交付", level=1)
doc.add_paragraph("一句话定位：念念 AI 是面向个人创作者和小型团队的项目级 AI 短剧创作平台，把剧本、参考素材、角色、镜头、真实任务和交付成果放进同一个可恢复项目。")
doc.add_paragraph("真实交付定义：用户完成“登录 → 新建项目 → 工作台 → 一键短剧 / 一键转绘 / 无限画布 → 发起真实任务 → 项目交付页 → 打开或下载真实 Word / 视频成果”。截图、HTTP 200、构建成功、Provider 任务号都只能证明中间步骤，不能替代用户可打开或下载的真实成果。")

doc.add_heading("2. 当前事实卡", level=1)
add_table(doc, ["事实项", "当前证据", "结论"], [
    ("线上 URL", "浏览器只读回读 https://ai.cauai.fun/，标题为“念念 AI - AI短剧创作平台”", "已确认"),
    ("线上首页", "线上 DOM 有首页/项目管理/工作台/导演台、开始创作、登录/注册；首页截图在本次浏览器回读中可见", "已确认"),
    ("线上 /#/projects 与 /#/workbench", "当前未登录回读均回落首页", "行为已观察；项目内登录态页面待核实"),
    ("线上 Studio 项目库", "线上 DOM 有新建空白项目、打开已有文件夹、60 秒预览、文本模型未接入提示、项目搜索和最近项目", "已确认"),
    ("线上导演台", "线上 DOM 有 3D 导演台区域、返回工作台按钮和 iframe", "入口已确认；内层 3D 细节待核实"),
    ("正式源码", "AGENTS.md 指定 E:\\codex\\niannianai\\niannianai 为独立 canonical root；release_baseline_review_evidence_20260809_canonical_main.json 记录 source_root 与 Git revision", "已确认"),
    ("本地共享文件基线", "同一 JSON 记录 server.js、index.html、app.js、CSS、sw.js 等 SHA-256；online_parity_claimed=false", "本地基线已确认，线上一致性未确认"),
    ("历史线上 release", "release-baselines/ai.cauai.fun/20260804...json 记录旧活动 release、回滚 release 和候选规则", "归档证据；不可直接作为当前源码"),
    ("已批准首页截图", "要求路径 authority/approved/homepage/r11-1440.png 在当前仓库不存在", "待核实，不能声称已读取"),
    ("V1 权威文档", "要求路径 documents/authoritative-docs/niannian-ai-web/AUTHORITATIVE_PROJECT_DOC_V1.md 在当前仓库不存在", "待核实，不能声称已读取"),
    ("authority/README 与工程 authority", "当前仓库无 authority/ 目录", "待补齐；本 V2 不能覆盖缺失证据"),
], widths=[1.45, 4.0, 1.5])
doc.add_paragraph("不可改范围（来自 AGENTS.md 与当前工程合同）：不恢复旧 Nomi 自建画布、#canvas 或 owned-canvas-director-import；保留当前念念首页、Logo、Studio 画布、素材库、Image2、H3、导演台、Step01-Step04 与已有 API 适配；不把密钥、Cookie、用户媒体或 Provider 原始响应写入源码、文档或 Git。")

doc.add_heading("3. 同类产品对标矩阵", level=1)
add_table(doc, ["对标行为", "念念 AI 具体页面 / 组件", "用户得到的结果", "不复制的边界"], [
    ("OiiOii：把 Agent、Style、Asset、Template 组织成用户可感知的创作过程", "工作台三入口 + 项目内阶段条；一键短剧显示剧本、角色、分镜、生成、交付状态", "用户知道当前项目处于哪个阶段，下一步缺什么，完成后得到什么", "不复制其导航、技能卡瀑布、图片和品牌文案；不把模型名变成主导航"),
    ("LibTV：Skill/模板货架与人类、Agent 两种入口", "工作台保留一键短剧、一键转绘、无限画布三个真实入口；Studio 资源区保留提示词库/技能库", "用户可从明确任务进入，也可在项目内补充技能和模板", "不复制其内容社区、模型市场和促销作品墙；入口服务念念项目，不服务营销"),
    ("RunningHub：可见任务流、画质/成本/阻塞状态", "项目交付页的业务状态条、输入证据、质量检查、可打开/可下载", "用户能判断是否完成、哪里阻塞、是否可交付", "不展示 Provider 编号、内部路径、密钥、签名地址或渠道术语；成本只呈现用户可理解的额度/费用提示"),
], widths=[2.0, 2.0, 1.7, 1.7])
doc.add_paragraph("差异化结论：念念以项目和真实交付为主语；模型、渠道和工作流只是必要的实现层。所有同类产品观察只改变信息结构原则，不复制布局、文案、资产或 Provider 分类。")

doc.add_heading("4. 前端理想图与信息结构", level=1)
doc.add_paragraph("理想图路径：documents/authoritative-docs/visuals/niannian-ai-web_ideal-home.png。当前文件是现有真实品牌素材的参考副本，因为 Krill 默认渠道 403、meinianda 渠道 429；它不是本次 Image2 生成结果，待渠道恢复后原位替换。完整提示词、参考路径、失败状态和人工审图结论见同目录 niannian-ai-web_ideal-home_prompt.md。")
if IDEAL.exists():
    doc.add_picture(str(IDEAL), width=Inches(6.4))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("图 1：目标首页/工作台视觉方向参考（非交互页面；非本次 Image2 生成结果）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

doc.add_heading("URL 树（当前产品表面）", level=2)
add_code(doc, "/\n├─ #projects                 主站项目管理（未登录回落首页，登录态待核实）\n├─ #workbench                主站工作台：四入口（无限画布/一键转绘/一键短剧/智能剪辑）\n├─ #director-desk            3D 导演台与返回工作台\n├─ studio/#/studio           Studio 项目库\n└─ studio/?step=...#/studio  Studio 画布、资源面板、生成检查器与时间轴")

doc.add_heading("三张核心页面 ASCII 线框", level=2)
add_code(doc, "[首页/工作台首屏]\n┌────────────────────────────────────────────────────────────┐\n│ 念念 AI   首页  项目管理  工作台  导演台       登录          │\n├───────────────────────┬────────────────────────────────────┤\n│ 当前项目 / 新建项目    │ 生产状态：剧本✓ 角色✓ 镜头生成中    │\n│ [开始新项目]            │ 交付：待确认 · 可打开 / 下载          │\n│ [一键短剧] [一键转绘]  │ 当前镜头 / 下一动作                  │\n│ [无限画布]              │ [继续制作]                         │\n└───────────────────────┴────────────────────────────────────┘")
add_code(doc, "[Studio 画布]\n┌──────┬───────────────────────────────────────┬──────────────┐\n│资源库 │ 项目 / 镜头   创作 生成 预览  保存中    │生成检查器     │\n│分组   │  [文本]──>[角色]──>[镜头]──>[视频]     │输入完整性 ✓   │\n│提示词 │        迷你地图 / 缩放 / 整理           │状态 / 规格     │\n│技能库 │ 底部时间轴：镜头 01 | 镜头 02 | 交付      │[生成][采用]   │\n└──────┴───────────────────────────────────────┴──────────────┘")
add_code(doc, "[项目交付页]\n┌────────────────────────────────────────────────────────────┐\n│ 项目名 / 一键短剧 / 交付                                    │\n├──────────────┬─────────────────────────────────────────────┤\n│ 阶段条         │ 成果预览                                    │\n│ 剧本 ✓         │ [真实 Word / 视频播放器]                    │\n│ 角色 ✓         │ 输入证据：项目、镜头、采用状态                 │\n│ 分镜 ✓         │ 质量检查：可读取 / 可播放                     │\n│ 交付 待确认    │ [打开成果] [下载成果] [重试失败步骤]           │\n└──────────────┴─────────────────────────────────────────────┘")

doc.add_heading("桌面与 390px 首屏验收点", level=2)
add_table(doc, ["视口", "必须一眼看见", "禁止"], [
    ("1440x900", "项目身份、创建主行动、三个真实生产入口、状态条、交付去向；共享导航和主 CTA 不被遮挡", "整页图片替代 UI、虚构统计大屏、紫色渐变、卡片嵌套"),
    ("390x844", "项目身份、阶段、真实状态、结果/输入证据、唯一主动作；无横向滚动", "缩小版自由画布、被软键盘遮挡、按钮小于 44px、长文本覆盖"),
], widths=[1.0, 4.5, 1.7])

doc.add_heading("5. 用户路径 DAG", level=1)
add_code(doc, "用户\n  ↓\n登录（真实 session）\n  ↓\n新建项目（项目 ID 写入服务端）\n  ↓\n工作台（选择项目生产入口）\n  ├─ 一键短剧：剧本/Word → 角色 → 分镜 → 真实生产 → Word 交付\n  ├─ 一键转绘：原片/参考素材 → Step01-Step04 → 真实视频交付\n  └─ 无限画布：节点/素材/提示词 → Image2/H3/文本任务 → 项目交付\n  ↓\n项目交付页（业务状态、输入证据、质量检查）\n  ↓\n打开或下载真实 Word / 视频成果")

doc.add_heading("6. 念念 AI 专属阶段路线", level=1)
add_stage(doc, "S0 权威基线与事实卡", "把线上、批准、正式源码、候选、归档和待核实项分开，冻结事实边界。", "线上只读 DOM/截图；AGENTS.md；PRODUCT.md；DESIGN.md；FRONTEND_EXECUTION_PLAN.md；release_baseline_review_evidence_20260809_canonical_main.json。", "一张可审阅事实卡和来源清单。", "只读回读线上核心路由；核对本地 canonical root、Git revision 和 SHA 记录；缺失 authority/V1 文件明确列为待核实。", "website-product-router（Source truth / existing repo redesign）", "website-quality-router（Workflow 1-2）", "主控 Agent 负责事实来源、阶段准入；任务执行 Agent 只读盘点，不改线上；同一文件范围单一写入者。", "不得用线上镜像覆盖正式源码；不得把旧候选或旧 Nomi 画布当成当前权威。", "每个“已实现”都有文件、线上回读或截图证据；待核实项不写成完成。", "事实卡和批准资产路径（若补齐）", "读取 AGENTS.md 与所有可用 authority/V1 证据，建立已确认/待核实表；不要改线上；输出来源、边界和下一阶段输入。")
add_stage(doc, "S1 视觉蓝图与可点击静态原型", "把事实卡转成可看见、可点击、但不接真实业务的首页/工作台蓝图。", "S0 事实卡；现有品牌主资源；三站公开页面观察；批准首页截图（当前缺失，待补齐）。", "理想图、三张 ASCII 线框、可点击静态原型和桌面/390px 验收记录。", "用 krill-image2 生成 bounded hero/方向图；用真实 HTML 组件实现按钮、入口和状态条；不把整图当网页。", "design-taste-frontend（Brief inference / redesign-preserve）", "krill-image2（Run / Verify）", "视觉原型 Agent 只写 visuals、线框和静态原型；不得改线上或真实 API。", "S1 只允许图、Word、线框、静态原型；不提交真实 Provider 任务，不写密钥。", "桌面和 390px 均能识别三个入口、创建行动、生产状态和交付方向；图像缺失时 HTML 仍可用。", "原型页面和审阅结论", "基于事实卡做视觉蓝图，使用真实品牌资源，生成方向图但保持 HTML-first；输出可点击静态原型和差异记录。")
add_stage(doc, "S2 项目中心与工作台骨架", "让登录后的用户能新建项目并进入四入口工作台，项目状态可恢复。", "S1 审阅通过的结构；真实 session、项目 API、当前首页/工作台路由。", "可点击前端骨架：登录 → 新建项目 → 工作台 → 入口选择。", "先实现项目实体和路由绑定，再接四入口卡片、项目状态条、继续动作、空/加载/失败状态；不接 Provider 生成。", "website-product-router（backend/data/admin route）", "website-quality-router（function/responsive gates）", "任务执行 Agent 是唯一写入者；主控 Agent 决定阶段准入；独立验证 Agent 只在候选冻结后检查，不改候选。", "不改变当前首页、Logo、Studio 画布和导演台；不把静态成功态当真实项目。", "清洁浏览器登录后新建项目，刷新仍回到同一项目；桌面/390px 入口可达。", "真实项目 ID 与工作台候选", "只在 S2 实现真实项目骨架；验证登录、新建、刷新和四入口路由，不触发 Provider 任务。")
add_stage(doc, "S3 一键短剧真实闭环：剧本/Word 输入 → 真实生产 → 第四步 Word 交付", "从剧本或 Word 输入连续产出可打开/下载的真实 Word 交付成果。", "S2 项目；真实 Word/剧本上传；现有 Step03/Step04 业务合同；文本模型配置。", "项目交付页上的可打开、可下载真实 Word。", "上传并持久化源文件；解析剧本；生成角色/分镜业务状态；执行已批准生产链；写入 Step04 交付；交付页只显示业务状态和恢复动作。", "mx-shortdrama-production-harness（启动/恢复/交付）", "mx-shortdrama-00-router（唯一专业路由）", "短剧生产 Agent 只在 S3 进入真实生产合同；页面只显示业务状态；主控负责费用和发布边界。", "不得显示 Provider 编号、密钥、原始响应；不得无故要求重新上传；不能用 mock Word 冒充成果。", "清洁浏览器从 Word 输入到第四步，用户能打开并下载真实 Word；中途刷新/离开后状态恢复。", "真实 Word 成果与交付记录", "使用真实项目和 Word 输入跑一条短剧生产链，验证第四步交付可打开/下载；失败必须保留输入并给出恢复动作。")
add_stage(doc, "S4 一键转绘真实闭环：原片/参考素材 → 真实视频交付", "从授权原片完成 Step01-Step04，并在项目交付页得到可播放/下载视频。", "S2 项目；授权原片；现有 Step01-Step04、Image2、H3 适配；真实服务器执行环境。", "项目交付页上的真实视频成果和采用记录。", "先做不付费合同/恢复/回读测试，再按既有授权做最短 5 秒试单；保存源素材、镜头、任务状态、结果引用和下载路径。", "mx-shortdrama-production-harness（启动/恢复/交付）", "image2-storyboard-video（Step04/视频结果）", "短剧生产 Agent 负责已批准生产链；独立验证 Agent 只验证候选，不改服务端。", "不切换 Provider/模型/路由；不超过 H3 单次 5 秒、总计 10000 秒；不把排队或任务号叫做交付。", "用户能播放、采用、下载真实视频；失败后可恢复且无需无故重传。", "真实视频交付与证据索引", "从已授权原片开始，跑最短真实镜头闭环；先完成合同和恢复测试，再做一次 5 秒试单；最终必须回读可播放视频。")
add_stage(doc, "S5 Nomi 无限画布与导演台整合", "保留当前念念 Studio 画布与导演台，把脚本、素材、节点、镜头和交付状态连成同一项目。", "S2 项目；Studio 画布现有节点/资源/API；导演台现有场景保存、截图和镜头合同。", "同一项目内可从画布/导演台组织镜头并进入交付页。", "先建立项目实体引用和素材使用关系；保留三模式、四类资源、八类节点、时间轴、生成侧栏和导演台返回路径；仅接已有真实 API。", "website-product-router（existing project redesign / source repo reuse）", "website-quality-router（responsive/function gates）", "任务执行 Agent 负责 Studio 写入；导演台验证 Agent 只在候选冻结后检查 3D/桌面/390px。", "不恢复旧 Nomi 自建画布、#canvas 或 owned-canvas-director-import；无服务端路径的音频/3D/白板/全景/场景 3D 只标编辑/参考。", "画布保存刷新、素材引用、导演台保存/返回、交付入口全部指向同一项目事实。", "统一项目图谱和交付投影", "在当前念念画布基线上整合导演台引用，验证保存/刷新/返回和项目交付投影；不替换现有画布家族。")
add_stage(doc, "S6 发布回读、线上基线固化与回滚准备", "仅在所有真实用户路径通过后，形成一个完整、可回读、可回滚的发布候选。", "S3-S5 的完整候选；桌面/390px 证据；线上基线；发布授权。", "版本化完整发布包、线上回读记录、回滚包和下一版 baseline。", "冻结单一候选；运行 CI、浏览器、任务恢复和交付验证；在明确授权后发布；回读线上 HTML/资产/关键路径并保存 hash；失败立即使用回滚包。", "website-quality-router（deployment gate / existing-site candidate discipline）", "browser/control-in-app-browser（线上回读）", "主控 Agent 负责发布决定和回滚决定；任务执行 Agent 只发布获准候选；独立验证 Agent 只验收不改包。", "S6 之前不得发布；不能把本地预览、截图、HTTP 200 或构建成功当作交付；不得从旧候选拼包。", "线上真实路径可复现；Word/视频成果可打开/下载；桌面/390px 与 baseline 无未解释回归；回滚包可用。", "下一版线上 baseline 与交接文档", "只发布已批准单一候选；完成线上 HTML/资产/用户路径回读、桌面/390px 比较和回滚演练；报告真实交付证据。")

doc.add_heading("7. Agent 与 AGENTS.md 最小协作合同", level=1)
add_table(doc, ["角色", "只负责", "明确禁止"], [
    ("主控 Agent", "用户决定、事实来源、阶段准入、费用/发布边界、最终交接", "不绕过事实证据，不把中间状态称交付"),
    ("任务管理与执行 Agent", "唯一当前写入者；候选实现、测试、真实回读", "同一范围并行写入；从旧候选覆盖线上"),
    ("视觉原型 Agent", "S1 图、线框、静态原型和审图", "接真实业务、改线上、保存密钥"),
    ("短剧生产 Agent", "仅 S3/S4 进入真实生产合同，页面只显示业务状态", "泄露 Provider 编号/原始响应；自行改产品范围"),
    ("独立验证 Agent", "候选冻结后检查浏览器、桌面/390px 和交付证据", "修改候选、绕过失败、直接发布"),
], widths=[1.5, 3.1, 3.1])
doc.add_paragraph("固定规则：同一写入范围只能有一个写入者；候选必须声明 parent baseline、范围和受保护表面；不得从旧候选覆盖线上；不得把旧自建画布恢复为正式画布；所有秘密只在受保护运行环境，不进入 GitHub、Word、截图或日志。")

doc.add_heading("8. 发布边界与验收证据", level=1)
add_table(doc, ["阶段", "允许", "不允许"], [
    ("S1", "图、Word、线框、静态原型", "真实 Provider、线上改动、发布"),
    ("S2", "可点击前端骨架、真实项目/session 接线", "假成功、Provider 生成"),
    ("S3-S5", "真实业务、真实任务、真实 Word/视频交付", "改变已批准 Provider/模型/路由；泄露内部细节"),
    ("S6", "单一候选发布、线上回读、回滚准备", "未授权发布、拼接旧候选、无回读交付"),
], widths=[1.0, 3.1, 3.6])
doc.add_paragraph("每次可见页面改动必须与线上 baseline 在 1440x900 与 390x844 比较。Word 生成完成后必须使用 documents Skill 的 render_docx.py 渲染为 PNG，逐页检查裁切、重叠、表格断裂、字体替换和页眉页脚；任何问题修复后重新渲染。")

doc.add_heading("9. 当前待我决定事项与下一动作", level=1)
add_bullets(doc, [
    "补齐并确认 authority/README、NIANNIAN_AI_ENGINEERING_AUTHORITY.md、online-baselines、approved 与 V1 文档，替换本 V2 中相应“待核实”标记。",
    "提供或恢复批准首页截图 authority/approved/homepage/r11-1440.png；在 Krill 渠道恢复后重新生成并审图理想首屏。",
    "决定 S0/S1 是否先冻结为 Word + 视觉方向图 + 静态原型，不进入真实业务接线。",
    "下一最小动作：补齐唯一 authority 目录和批准截图，然后对 V2 做一次事实卡复核；在此之前不改线上。",
])

doc.add_heading("附录 A：来源与证据索引", level=1)
add_bullets(doc, [
    "项目合同：AGENTS.md",
    "产品与界面合同：PRODUCT.md、DESIGN.md、FRONTEND_EXECUTION_PLAN.md",
    "同类产品研究：docs/frontend-reference-study.md",
    "原站继承矩阵：docs/frontend-inheritance-matrix.md",
    "本地源码基线：release_baseline_review_evidence_20260809_canonical_main.json、release_baseline_attestation_20260809_canonical_main.json",
    "历史线上 release 证据：release-baselines/ai.cauai.fun/20260804-workbench-clarity-r2-short-drama-modal-fix1.json",
    "线上只读回读：https://ai.cauai.fun/、/#/projects、/#/workbench、/studio/#/studio、/#director-desk（2026-08-10）",
    "视觉参考：assets/home/niannian-hero-oil-paint-quiet-v1.png；Image2 提示词与失败记录：documents/authoritative-docs/visuals/niannian-ai-web_ideal-home_prompt.md",
])

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("念念 AI 主站权威项目文档 V2 · 内部前期基线 · 不代表线上发布").font.size = Pt(8)

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(OUT)
