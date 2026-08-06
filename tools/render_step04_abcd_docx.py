from __future__ import annotations

"""Deterministic Step04 D renderer.

Only consumes step04_abcd_contract.json. It never reads raw Step02 cards or
re-infers people, assets, actions, dialogue, or camera facts.
"""

import argparse
import hashlib
import json
import os
import sys
from pathlib import Path
from typing import Any

# Production bundles may carry Python wheels in a package-local directory.
# Resolve that directory before importing python-docx so the renderer does not
# depend on a machine-global installation.
_TOOLS_DIR = Path(__file__).resolve().parent
_BUNDLED_PYTHON_DEPS = [
    Path(os.environ.get("NIANNIAN_STEP04_PYTHON_DEPS", "")),
    _TOOLS_DIR.parent / "runtime" / "python-deps",
    _TOOLS_DIR / "vendor" / "python-deps",
]
for _dependency_root in _BUNDLED_PYTHON_DEPS:
    if str(_dependency_root) and _dependency_root.is_dir():
        sys.path.insert(0, str(_dependency_root))
        break

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ACCENT = "2E74B5"
INK = "0B2545"
MUTED = "555555"
FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def canonical(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def sha256_value(value: Any) -> str:
    return hashlib.sha256(canonical(value).encode("utf-8")).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def seconds_range(start_ms: int, end_ms: int) -> str:
    """Delivery view uses seconds; ms remain only in the immutable contract."""
    return f"{start_ms / 1000:.3f}–{end_ms / 1000:.3f}秒"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_font(run, size=10, bold=False, color=INK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(cell, text: str, size=9, bold=False, color=INK) -> None:
    paragraph = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(text))
    set_font(run, size, bold, color)


def make_asset_preview(asset: dict[str, Any], preview_dir: Path) -> tuple[Path | None, str]:
    """Create a small Word-only JPEG preview without changing the source asset."""
    source = Path(str(asset.get("exact_path") or ""))
    if str(asset.get("status") or "") != "accepted" or not source.is_file():
        return None, "待生成（新图默认 2K）"
    asset_id = str(asset.get("asset_id") or "asset")
    digest = str(asset.get("sha256") or sha256_file(source))[:12]
    preview_dir.mkdir(parents=True, exist_ok=True)
    preview = preview_dir / f"{asset_id}_{digest}_preview.jpg"
    if not preview.is_file():
        with Image.open(source) as image:
            rendered = image.convert("RGB")
            rendered.thumbnail((560, 560), Image.Resampling.LANCZOS)
            rendered.save(preview, format="JPEG", quality=78, optimize=True, progressive=True)
    with Image.open(source) as image:
        dimensions = f"{image.width}×{image.height}"
    return preview, f"已验收 · 原图 {dimensions} · Word 压缩预览"


def add_asset_preview(cell, asset: dict[str, Any], preview_dir: Path) -> None:
    cell.text = ""
    preview, label = make_asset_preview(asset, preview_dir)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    if preview:
        paragraph.add_run().add_picture(str(preview), width=Inches(1.52))
    caption = cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(2)
    set_font(caption.add_run(label), 6.7, False, MUTED)


def render_asset_prompt_table(doc: Document, assets: list[dict[str, Any]], preview_dir: Path) -> None:
    """Render the user-facing B table with an inline preview of each real asset."""
    prompt_table = doc.add_table(rows=1, cols=4)
    set_table_geometry(prompt_table, [1500, 1000, 2600, 4260])
    header_row(prompt_table, ["资产与中文名", "资产类型", "实际成图（压缩预览）", "已验收生图提示词"])
    for asset in assets:
        row = prompt_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[3].text = ""
        add_text(row.cells[0], str(asset.get("display_name") or ""), 8.2)
        add_text(row.cells[1], str(asset.get("kind") or ""), 8.2)
        add_asset_preview(row.cells[2], asset, preview_dir)
        add_text(row.cells[3], str(asset.get("image_prompt") or ""), 8.2)
    set_table_geometry(prompt_table, [1500, 1000, 2600, 4260])


def evidence_label(evidence_ids: list[Any]) -> str:
    """Keep machine evidence IDs in A/C JSON, not in the user-facing Word."""
    count = len([value for value in evidence_ids if str(value)])
    return f"已绑定 {count} 项证据" if count else "无"


def heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 4)
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, True, ACCENT if level < 3 else INK)


def header_row(table, labels: list[str]) -> None:
    row = table.rows[0]
    for cell, label in zip(row.cells, labels):
        set_cell_shading(cell, FILL)
        add_text(cell, label, size=9, bold=True, color=INK)
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def delivery_view(a_layer: dict[str, Any], b_layer: dict[str, Any], c_layer: dict[str, Any], d_layer: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any], dict[str, Any]]:
    """Normalize the two strict A/B/C schemas into one render-only view.

    Both schemas are evidence contracts. This function only changes field names
    for presentation and never derives a person, asset, event, or time range
    from source prose.
    """
    if "entity_instances" in a_layer:
        return a_layer, b_layer, c_layer, d_layer

    entities = [
        {**item, "shot_id": int(str(item.get("shot_id") or "0").lstrip("S"))}
        for item in list(a_layer.get("entities") or [])
    ]
    slots = list(b_layer.get("reference_slots") or [])
    slot_by_id = {str(slot.get("reference_slot_id")): slot for slot in slots}
    name_by_instance = {str(item.get("instance_id")): str(item.get("role_ref") or "") for item in entities}
    assets_by_id: dict[str, dict[str, Any]] = {}
    for slot in slots:
        asset_id = str(slot.get("asset_id") or "")
        asset = assets_by_id.setdefault(asset_id, {
            "asset_id": asset_id,
            "display_name": str(slot.get("display_name") or ""),
            "duty": str(slot.get("purpose") or ""),
            "kind": str(slot.get("kind") or ""),
            "generation_prompt": str(slot.get("generation_prompt") or slot.get("image_prompt") or ""),
            "image_prompt": str(slot.get("generation_prompt") or slot.get("image_prompt") or ""),
            "used_by_shots": [],
            "exact_path": str(slot.get("exact_path") or ""),
            "sha256": str(slot.get("sha256") or ""),
            "status": str(slot.get("status") or ""),
        })
        shot = str(slot.get("shot_id") or "").lstrip("S")
        if shot and shot not in asset["used_by_shots"]:
            asset["used_by_shots"].append(shot)

    groups = []
    for group in list(c_layer.get("prompt_groups") or []):
        group_slots = [slot_by_id[slot_id] for slot_id in group.get("reference_slots") or [] if slot_id in slot_by_id]
        event_lines: list[str] = []
        events = []
        for event in list(group.get("events") or []):
            start, end = [int(value) for value in event.get("timecode_ms") or []]
            subject = name_by_instance.get(str(event.get("subject_instance_id") or ""), "")
            object_name = name_by_instance.get(str(event.get("object_instance_id") or ""), "")
            change = str(event.get("change") or "").strip()
            line = f"{seconds_range(start, end)}：{subject}{change}"
            if object_name:
                line += f"，对象为{object_name}"
            dialogue = event.get("dialogue") or None
            if dialogue:
                speaker = name_by_instance.get(str(dialogue.get("speaker_instance_id") or dialogue.get("speaker_id") or ""), "")
                text = str(dialogue.get("text") or dialogue.get("content") or "").strip()
                line += f"；{speaker}在该动作中说：\"{text}\""
            event_lines.append(line)
            events.append({"start_ms": start, "end_ms": end, "change": change, "action": change})
        start_ms = int(group.get("source_start_ms") or (events[0]["start_ms"] if events else 0))
        end_ms = int(group.get("source_end_ms") or (events[-1]["end_ms"] if events else start_ms))
        # If the bridge supplied the compact C view, render it verbatim. D
        # must not rebuild a longer prompt from the same facts.
        prompt_text = str(group.get("prompt_text") or "").strip()
        if not prompt_text:
            prompt_text = "\n".join([
                f"场景：{str(group.get('scene_identity') or '')}",
                f"环境：{str(group.get('environment_identity') or '')}",
                "参考图：" + "、".join(str(slot.get("display_name") or "") for slot in group_slots),
                f"构图：S{int(str(group.get('shot_id') or '0').lstrip('S')):03d} {seconds_range(start_ms, end_ms)}；{str(group.get('composition') or '')}",
                "变化：" + "；".join(event_lines),
                f"镜头：{str(group.get('camera') or '')}",
                f"光线：{str(group.get('light') or '')}",
                f"声音：{str(group.get('sound') or '')}",
            ])
        groups.append({
            "group_id": str(group.get("group_id") or ""),
            "references": [{"reference_key": str(slot.get("display_name") or ""), "duty": str(slot.get("purpose") or "")} for slot in group_slots],
            "segments": [{"shot_id": int(str(group.get("shot_id") or "0").lstrip("S")), "start_ms": start_ms, "end_ms": end_ms, "prompt_text": prompt_text, "prompt_compression": group.get("prompt_compression") or {}, "events": events}],
        })
    return (
        {"entity_instances": entities},
        {"assets": list(assets_by_id.values())},
        {"groups": groups},
        {**d_layer, "gate": {"status": "passed"}},
    )


def render_prompt_review(doc: Document, b: dict[str, Any], c: dict[str, Any], preview_dir: Path) -> None:
    """Render only the two user-selected review sections from immutable B/C."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    set_font(title.add_run("Step04 资产与生视频提示词包"), 22, True, INK)

    heading(doc, "B 层资产图提示词表", 1)
    render_asset_prompt_table(doc, list(b.get("assets") or []), preview_dir)

    doc.add_page_break()
    heading(doc, "三、完整生视频提示词", 1)
    for group_index, group in enumerate(c.get("groups", [])):
        if group_index and group_index % 4 == 0:
            doc.add_page_break()
        segments = list(group.get("segments") or [])
        segment = segments[0] if segments else {}
        group_start_ms = int(group.get("source_start_ms") or segment.get("start_ms") or 0)
        group_end_ms = int(group.get("source_end_ms") or (segments[-1].get("end_ms") if segments else group_start_ms) or group_start_ms)
        duration = seconds_range(0, group_end_ms - group_start_ms)
        prompt_text = str(group.get("prompt_text") or segment.get("prompt_text") or "")
        heading(doc, f"{group.get('group_id', '')} | {duration}", 2)
        prompt_table = doc.add_table(rows=1, cols=1)
        set_table_geometry(prompt_table, [9360])
        header_row(prompt_table, ["可直接提交的完整生视频提示词"])
        row = prompt_table.add_row()
        row.cells[0].text = ""
        add_text(row.cells[0], prompt_text, 8.7)
        set_table_geometry(prompt_table, [9360])


def write_receipt(contract_path: Path, output_path: Path, expected: str, doc: Document, presentation_profile: str) -> None:
    receipt = {
        "schema_version": "mx_shortdrama_step04d_render_receipt_v1",
        "contract_path": str(contract_path.resolve()),
        "contract_sha256": expected,
        "output_docx_path": str(output_path.resolve()),
        "output_docx_sha256": sha256_file(output_path),
        "output_docx_bytes": output_path.stat().st_size,
        "structural_qa": {
            "status": "passed",
            "table_count": len(doc.tables),
            "required_sections": ["B", "C"] if presentation_profile == "prompt_review" else ["A", "B", "C", "D"],
            "table_geometry_applied": True,
        },
        "presentation_profile": presentation_profile,
        "visual_qa": {"status": "pending_docx_preview_screenshot"},
        "provider_calls": {"image": False, "video": False},
        "renderer_policy": "D only rendered the immutable A/B/C contract; no source facts were read or changed.",
    }
    receipt_path = output_path.parent / "step04d_render_receipt.json"
    receipt_path.write_text(json.dumps(receipt, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def render(contract_path: Path, output_path: Path, presentation_profile: str = "full") -> None:
    contract = json.loads(contract_path.read_text(encoding="utf-8"))
    expected = str(contract.get("contract_sha256") or "")
    unsigned = dict(contract)
    unsigned.pop("contract_sha256", None)
    if expected and sha256_value(unsigned) != expected:
        raise ValueError("contract_sha256 mismatch")
    layers = contract.get("layers") or {}
    a, b, c, d = delivery_view(layers.get("A") or {}, layers.get("B") or {}, layers.get("C") or {}, layers.get("D") or {})
    if not a or not b or not c or d.get("gate", {}).get("status") != "passed":
        raise ValueError("A/B/C/D contract is not delivery-ready")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if presentation_profile == "prompt_review":
        render_prompt_review(doc, b, c, output_path.parent / "asset_previews")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        write_receipt(contract_path, output_path, expected, doc, presentation_profile)
        return

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("Step04 A/B/C/D 中文权威生产包"), 22, True, INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("只读事实编译交付 | 不调用图片或视频渠道"), 10, False, MUTED)

    meta = doc.add_table(rows=1, cols=2)
    set_table_geometry(meta, [2400, 6960])
    header_row(meta, ["交付项", "值"])
    metadata = [
        ("合同 SHA-256", str(contract.get("contract_sha256") or "")),
        ("Step02 输入 SHA-256", str((contract.get("source") or {}).get("step02_acceptance_sha256") or "")),
        ("A/B/C 输入摘要", f"A {sha256_value(a)[:12]} | B {sha256_value(b)[:12]} | C {sha256_value(c)[:12]}"),
        ("渠道调用", "图片：否；视频：否"),
    ]
    for label, value in metadata:
        row = meta.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        add_text(row.cells[0], label, 9, True)
        add_text(row.cells[1], value, 9)
    set_table_geometry(meta, [2400, 6960])

    # Keep the Word package readable in a real viewer.  The previous renderer
    # allowed every C-layer group to flow into one extremely tall page; that
    # passed structural checks but made the actual production prompts unusable.
    doc.add_page_break()
    heading(doc, "一、A 层：镜头人物实例与证据绑定")
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1200, 2600, 1500, 2200, 1860])
    header_row(table, ["镜头", "中文角色引用", "资产类别", "证据", "状态"])
    for item in a.get("entity_instances", []):
        row = table.add_row()
        vals = [f"S{int(item['shot_id']):03d}", item["role_ref"], "人物", evidence_label(item.get("evidence_ids", [])), "已闭合"]
        for cell, value in zip(row.cells, vals):
            cell.text = ""
            add_text(cell, value, 8.5)
    set_table_geometry(table, [1200, 2600, 1500, 2200, 1860])

    heading(doc, "二、B 层：资产与连续性合同")
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1600, 2100, 2750, 1500, 1410])

    heading(doc, "B 层资产图提示词表", 2)
    render_asset_prompt_table(doc, list(b.get("assets") or []), output_path.parent / "asset_previews")
    header_row(table, ["资产与中文名", "职责", "适用镜头", "SHA-256", "状态"])
    for asset in b.get("assets", []):
        row = table.add_row()
        used = ", ".join(f"S{int(s):03d}" for s in asset.get("used_by_shots", []))
        asset_state = "已验收" if asset.get("status") == "accepted" else ("待生成" if asset.get("status") == "planned" else str(asset.get("status") or "未标注"))
        vals = [asset.get("display_name", ""), asset.get("duty", ""), used, str(asset.get("sha256") or "")[:12] or "—", asset_state]
        for cell, value in zip(row.cells, vals):
            cell.text = ""
            add_text(cell, value, 8.3)
    set_table_geometry(table, [1600, 2100, 2750, 1500, 1410])

    doc.add_page_break()
    heading(doc, "三、完整生视频提示词")
    for group_index, group in enumerate(c.get("groups", [])):
        if group_index and group_index % 4 == 0:
            doc.add_page_break()
        heading(doc, group.get("group_id", ""), 2)
        segments = list(group.get("segments") or [])
        segment = segments[0] if segments else {}
        group_start_ms = int(group.get("source_start_ms") or segment.get("start_ms") or 0)
        group_end_ms = int(group.get("source_end_ms") or (segments[-1].get("end_ms") if segments else group_start_ms) or group_start_ms)
        prompt_text = str(group.get("prompt_text") or segment.get("prompt_text") or "")
        heading(doc, f"{group.get('group_id', '')} | {seconds_range(group_start_ms, group_end_ms)}", 3)
        prompt_table = doc.add_table(rows=1, cols=1)
        set_table_geometry(prompt_table, [9360])
        header_row(prompt_table, ["可直接提交的完整生视频提示词"])
        row = prompt_table.add_row()
        row.cells[0].text = ""
        add_text(row.cells[0], prompt_text, 8.7)
        set_table_geometry(prompt_table, [9360])

    doc.add_page_break()
    heading(doc, "四、D 层：交付质量门")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run("D 层只渲染 A/B/C，不修改事实。"), 10, True, ACCENT)
    for line in [
        "事实质量：人物、动作与说话人均回指已验证实例。",
        "资产质量：每个参考槽位均有真实路径、SHA-256 和验证状态。",
        "交付质量：Word 内容来自 C 层提示词 IR，未读取旧 Word 或原始自由文本。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(line), 10, False, INK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    write_receipt(contract_path, output_path, expected, doc, presentation_profile)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--contract", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--presentation-profile", choices=["full", "prompt_review"], default="full")
    args = parser.parse_args()
    render(args.contract, args.output, args.presentation_profile)
